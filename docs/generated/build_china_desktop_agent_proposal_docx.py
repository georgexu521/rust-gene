from __future__ import annotations

import math
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "generated"
ASSET_DIR = ROOT / "docs" / "proposal_assets"
DOCX_PATH = OUT_DIR / "china_desktop_ai_project_studio_proposal_2026-05-21.docx"

ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUT_DIR.mkdir(parents=True, exist_ok=True)


NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "1B8A8F"
GOLD = "C58B2B"
INK = "172033"
MUTED = "667085"
LIGHT = "F4F6F9"
BORDER = "D9E2EC"
GREEN = "18794E"
RED = "B42318"
PAPER = "F8FAFC"


def rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def pil_rgb(hex_color: str) -> tuple[int, int, int]:
    hex_color = hex_color.lstrip("#")
    return (int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def font_path() -> str:
    candidates = [
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return ""


FONT_PATH = font_path()


def f(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if FONT_PATH:
        try:
            return ImageFont.truetype(FONT_PATH, size=size, index=0)
        except Exception:
            pass
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        current = ""
        for char in paragraph:
            candidate = current + char
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = char
        if current:
            lines.append(current)
    return lines or [""]


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str,
    spacing: int = 6,
) -> None:
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 28
    lines = wrap_text(draw, text, font, max_width)
    line_heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(line_heights) + spacing * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for i, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=pil_rgb(fill))
        y += line_heights[i] + spacing


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str, width: int = 5) -> None:
    draw.line([start, end], fill=pil_rgb(fill), width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head = 18
    points = [
        end,
        (int(end[0] - head * math.cos(angle - math.pi / 6)), int(end[1] - head * math.sin(angle - math.pi / 6))),
        (int(end[0] - head * math.cos(angle + math.pi / 6)), int(end[1] - head * math.sin(angle + math.pi / 6))),
    ]
    draw.polygon(points, fill=pil_rgb(fill))


def save_cover_image() -> Path:
    path = ASSET_DIR / "figure_00_cover_concept.png"
    img = Image.new("RGB", (1600, 900), pil_rgb(PAPER))
    draw = ImageDraw.Draw(img)

    # Background accents
    draw.rounded_rectangle((80, 80, 1520, 820), radius=42, fill=(255, 255, 255), outline=pil_rgb(BORDER), width=3)
    draw.rounded_rectangle((120, 140, 1480, 760), radius=28, fill=(246, 249, 252), outline=(229, 236, 244), width=2)

    # Desktop window
    draw.rounded_rectangle((180, 190, 1420, 710), radius=28, fill=(255, 255, 255), outline=(210, 222, 235), width=3)
    draw.rounded_rectangle((180, 190, 1420, 250), radius=28, fill=pil_rgb(NAVY))
    for i, c in enumerate(["FF6B6B", "FFD166", "06D6A0"]):
        draw.ellipse((220 + i * 38, 210, 242 + i * 38, 232), fill=pil_rgb(c))
    draw.text((380, 208), "Priority Studio - AI 项目工作台", font=f(28), fill=(255, 255, 255))

    # Chat column
    draw.rounded_rectangle((220, 285, 710, 675), radius=20, fill=(248, 250, 252), outline=(224, 232, 240))
    draw.text((250, 310), "用户只说需求", font=f(32), fill=pil_rgb(NAVY))
    bubbles = [
        ("帮我做一个咖啡店会员积分工具", (250, 365, 640, 425), BLUE),
        ("正在创建项目、安装依赖、启动预览", (285, 455, 665, 515), TEAL),
        ("预览已就绪，可继续修改", (250, 545, 610, 605), GREEN),
    ]
    for text, box, color in bubbles:
        draw.rounded_rectangle(box, radius=18, fill=pil_rgb(color), outline=pil_rgb(color))
        draw_centered_text(draw, box, text, f(24), "FFFFFF")

    # Preview column
    draw.rounded_rectangle((770, 285, 1375, 675), radius=20, fill=(255, 255, 255), outline=(224, 232, 240))
    draw.text((810, 310), "后台完成工程闭环", font=f(32), fill=pil_rgb(NAVY))
    steps = [
        ("创建文件", 815, 375, BLUE),
        ("运行验证", 1040, 375, TEAL),
        ("自动修复", 815, 500, GOLD),
        ("交付结果", 1040, 500, GREEN),
    ]
    for label, x, y, color in steps:
        draw.rounded_rectangle((x, y, x + 185, y + 82), radius=18, fill=pil_rgb(color), outline=pil_rgb(color))
        draw_centered_text(draw, (x, y, x + 185, y + 82), label, f(25), "FFFFFF")
    draw_arrow(draw, (1000, 416), (1032, 416), NAVY, 4)
    draw_arrow(draw, (910, 462), (910, 492), NAVY, 4)
    draw_arrow(draw, (1000, 541), (1032, 541), NAVY, 4)

    # Model chips
    y = 645
    x = 825
    for label in ["DeepSeek", "Kimi", "GLM", "Qwen", "MiniMax"]:
        w = 35 + len(label) * 17
        draw.rounded_rectangle((x, y, x + w, y + 42), radius=21, fill=(238, 244, 251), outline=(204, 219, 235))
        draw_centered_text(draw, (x, y, x + w, y + 42), label, f(20), DARK_BLUE)
        x += w + 12

    draw.text((120, 782), "封面概念图：不展示代码的桌面项目生产体验", font=f(24), fill=pil_rgb(MUTED))
    img.save(path, quality=95)
    return path


def save_positioning_image() -> Path:
    path = ASSET_DIR / "figure_01_market_positioning.png"
    img = Image.new("RGB", (1600, 960), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((80, 60), "市场定位：从开发者命令行，转向普通用户桌面项目工作台", font=f(42), fill=pil_rgb(NAVY))
    draw.text((80, 115), "横轴代表用户技术门槛，纵轴代表产品输出是否以可用结果为中心。", font=f(24), fill=pil_rgb(MUTED))

    left, top, right, bottom = 180, 190, 1420, 810
    draw.rectangle((left, top, right, bottom), outline=pil_rgb(BORDER), width=3)
    draw.line((left, (top + bottom) // 2, right, (top + bottom) // 2), fill=pil_rgb(BORDER), width=2)
    draw.line(((left + right) // 2, top, (left + right) // 2, bottom), fill=pil_rgb(BORDER), width=2)
    draw.text((left, bottom + 30), "低技术门槛", font=f(24), fill=pil_rgb(MUTED))
    draw.text((right - 145, bottom + 30), "高技术门槛", font=f(24), fill=pil_rgb(MUTED))
    draw.text((30, top + 10), "结果/项目中心", font=f(24), fill=pil_rgb(MUTED))
    draw.text((30, bottom - 35), "代码/工具中心", font=f(24), fill=pil_rgb(MUTED))

    quadrants = [
        ((left + 30, top + 30, (left + right) // 2 - 30, (top + bottom) // 2 - 30), "普通聊天 AI", "会回答，但不负责项目闭环", "F4F6F9"),
        (((left + right) // 2 + 30, top + 30, right - 30, (top + bottom) // 2 - 30), "专业 Agent / IDE", "强能力，但默认面向程序员", "F4F6F9"),
        ((left + 30, (top + bottom) // 2 + 30, (left + right) // 2 - 30, bottom - 30), "传统模板工具", "模板强，但智能与工程闭环弱", "F4F6F9"),
        (((left + right) // 2 + 30, (top + bottom) // 2 + 30, right - 30, bottom - 30), "CLI 编程工具", "需要命令行与工程知识", "F4F6F9"),
    ]
    for box, title, sub, fill in quadrants:
        draw.rounded_rectangle(box, radius=18, fill=pil_rgb(fill), outline=(231, 236, 242))
        draw_centered_text(draw, (box[0], box[1] + 24, box[2], box[1] + 88), title, f(25), MUTED)
        draw_centered_text(draw, (box[0] + 34, box[1] + 88, box[2] - 34, box[3] - 22), sub, f(21), MUTED)

    def point(x: int, y: int, label: str, color: str, w: int = 190) -> None:
        draw.rounded_rectangle((x, y, x + w, y + 58), radius=18, fill=pil_rgb(color), outline=pil_rgb(color))
        draw_centered_text(draw, (x, y, x + w, y + 58), label, f(22), "FFFFFF")

    point(1000, 615, "Claude Code / Codex", DARK_BLUE, 270)
    point(1040, 700, "Qwen Code / Crush", BLUE, 255)
    point(390, 610, "传统模板工具", GOLD, 220)
    point(315, 235, "Priority Studio", TEAL, 260)
    draw.text((305, 306), "目标区域：低门槛 + 项目可交付", font=f(24), fill=pil_rgb(TEAL))

    img.save(path, quality=95)
    return path


def save_loop_image() -> Path:
    path = ASSET_DIR / "figure_02_product_loop.png"
    img = Image.new("RGB", (1600, 900), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((80, 60), "产品闭环：用户只表达目标，Agent 负责工程执行和验证", font=f(42), fill=pil_rgb(NAVY))
    draw.text((80, 115), "面向普通用户的关键不是“生成代码”，而是持续交付可运行、可预览、可修改的项目。", font=f(24), fill=pil_rgb(MUTED))

    center = (800, 485)
    radius_x, radius_y = 520, 250
    steps = [
        ("1 需求澄清", "少量关键问题", BLUE),
        ("2 创建项目", "模板与工作区", TEAL),
        ("3 写入文件", "代码与资源", DARK_BLUE),
        ("4 安装依赖", "环境托管", GOLD),
        ("5 运行预览", "用户看结果", GREEN),
        ("6 自动验证", "命令与测试", BLUE),
        ("7 失败修复", "诊断与补丁", RED),
        ("8 交付/回退", "快照与导出", TEAL),
    ]
    points = []
    for i in range(len(steps)):
        angle = -math.pi / 2 + i * 2 * math.pi / len(steps)
        x = int(center[0] + radius_x * math.cos(angle))
        y = int(center[1] + radius_y * math.sin(angle))
        points.append((x, y))

    for i, start in enumerate(points):
        end = points[(i + 1) % len(points)]
        draw_arrow(draw, start, end, BORDER, 4)

    for (title, sub, color), (x, y) in zip(steps, points):
        box = (x - 140, y - 58, x + 140, y + 58)
        draw.rounded_rectangle(box, radius=18, fill=(255, 255, 255), outline=pil_rgb(color), width=4)
        draw_centered_text(draw, (box[0], box[1] + 8, box[2], box[1] + 50), title, f(24), color)
        draw_centered_text(draw, (box[0], box[1] + 48, box[2], box[3] - 8), sub, f(19), MUTED)

    draw.rounded_rectangle((570, 390, 1030, 580), radius=28, fill=pil_rgb(LIGHT), outline=pil_rgb(BORDER), width=2)
    draw_centered_text(draw, (600, 405, 1000, 500), "AI 项目工作台", f(38), NAVY)
    draw_centered_text(draw, (610, 500, 990, 560), "聊天前台 + 本地执行 runtime", f(24), MUTED)

    img.save(path, quality=95)
    return path


def save_architecture_image() -> Path:
    path = ASSET_DIR / "figure_03_architecture.png"
    img = Image.new("RGB", (1600, 980), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((80, 60), "建议架构：桌面壳 + 本地 Rust Agent Runtime + 国内多模型网关", font=f(42), fill=pil_rgb(NAVY))
    draw.text((80, 115), "把现有 priority-agent 能力沉到后台，用桌面前台隐藏代码、命令和复杂环境。", font=f(24), fill=pil_rgb(MUTED))

    layers = [
        ("桌面\n应用层", "Tauri / React 或 Vue\n聊天、项目列表、预览、导出、权限弹窗", BLUE),
        ("Agent\nRuntime 层", "会话、意图路由、工具调用、文件操作、终端进程、验证、修复、追踪、记忆", TEAL),
        ("项目\n执行层", "隔离工作区、依赖安装、开发服务器、快照、回滚、敏感信息扫描", DARK_BLUE),
        ("模型\n渠道层", "DeepSeek / Kimi / GLM / Qwen / MiniMax\n多模型路由、成本控制、企业账号", GOLD),
    ]
    y = 190
    for title, body, color in layers:
        draw.rounded_rectangle((130, y, 1470, y + 145), radius=24, fill=(248, 250, 252), outline=pil_rgb(color), width=4)
        draw.rounded_rectangle((130, y, 420, y + 145), radius=24, fill=pil_rgb(color), outline=pil_rgb(color))
        draw_centered_text(draw, (160, y + 18, 390, y + 127), title, f(28), "FFFFFF")
        draw_centered_text(draw, (455, y + 20, 1430, y + 125), body, f(25), INK)
        if y < 625:
            draw_arrow(draw, (800, y + 152), (800, y + 192), NAVY, 5)
        y += 185

    # Side evidence rail
    draw.rounded_rectangle((130, 885, 1470, 935), radius=16, fill=pil_rgb(LIGHT), outline=pil_rgb(BORDER))
    draw_centered_text(draw, (150, 890, 1450, 930), "关键护城河：本地项目上下文、验证证据、失败修复经验、用户工作流记忆，而不是绑定单一模型。", f(23), DARK_BLUE)
    img.save(path, quality=95)
    return path


def save_roadmap_image() -> Path:
    path = ASSET_DIR / "figure_04_roadmap.png"
    img = Image.new("RGB", (1600, 780), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((80, 60), "6 个月落地路线图", font=f(42), fill=pil_rgb(NAVY))
    draw.text((80, 115), "用真实项目完成率验证价值，而不是只做 Demo 或堆功能。", font=f(24), fill=pil_rgb(MUTED))

    x0, y0 = 150, 260
    width = 1300
    months = ["M1", "M2", "M3", "M4", "M5", "M6"]
    for i, month in enumerate(months):
        x = x0 + i * width // 6
        draw.line((x, y0 - 45, x, y0 + 345), fill=pil_rgb(BORDER), width=2)
        draw.text((x + 78, y0 - 85), month, font=f(25), fill=pil_rgb(MUTED))
    draw.line((x0, y0 + 345, x0 + width, y0 + 345), fill=pil_rgb(BORDER), width=3)

    bars = [
        ("阶段 1  桌面壳 + Runtime", 0, 2, BLUE, "Tauri、聊天、工作区、Web 闭环"),
        ("阶段 2  非程序员 MVP", 2, 5, TEAL, "预览、验证修复、快照、导出、国内模型"),
        ("阶段 3  内测验证", 5, 6, GOLD, "50-100 名用户\n成本与付费场景"),
    ]
    y = y0
    for title, start, end, color, detail in bars:
        bx1 = x0 + start * width // 6 + 10
        bx2 = x0 + end * width // 6 - 10
        draw.rounded_rectangle((bx1, y, bx2, y + 85), radius=18, fill=pil_rgb(color), outline=pil_rgb(color))
        title_font = f(24 if (bx2 - bx1) > 360 else 21)
        body_font = f(19 if (bx2 - bx1) > 360 else 17)
        draw.text((bx1 + 24, y + 12), title, font=title_font, fill=(255, 255, 255))
        detail_lines = detail.split("\n")
        for line_idx, line in enumerate(detail_lines):
            draw.text((bx1 + 24, y + 47 + line_idx * 24), line, font=body_font, fill=(255, 255, 255))
        y += 110

    metrics = [
        "首个可运行网页项目闭环",
        "20 个任务 ≥14 个完成",
        "找到 2-3 个高频付费场景",
    ]
    y = 625
    x = 170
    for metric in metrics:
        draw.rounded_rectangle((x, y, x + 390, y + 62), radius=18, fill=pil_rgb(LIGHT), outline=pil_rgb(BORDER))
        draw_centered_text(draw, (x + 15, y + 5, x + 375, y + 57), metric, f(22), DARK_BLUE)
        x += 430

    img.save(path, quality=95)
    return path


def save_progress_image() -> Path:
    path = ASSET_DIR / "figure_05_development_progress.png"
    img = Image.new("RGB", (1600, 900), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((80, 60), "当前开发进展：已有 Agent Runtime，正在转向桌面产品化", font=f(42), fill=pil_rgb(NAVY))
    draw.text((80, 115), "合作方需要看到：项目已有工程底座，但桌面应用、用户体验和商业闭环仍是下一阶段重点。", font=f(24), fill=pil_rgb(MUTED))

    columns = [
        ("已具备", "Rust Agent Runtime\n工具执行 / 权限 / 记忆\ntrace / required validation\n多 provider 协议基础", TEAL),
        ("验证中", "真实项目 gauntlet\n修复与 closeout 证据\nprovider 协议回归矩阵\n风险信号与工作流 targeting", BLUE),
        ("下一步", "Tauri 桌面应用\n非程序员 onboarding\n运行预览和导出\n国内模型套餐与内测", GOLD),
    ]
    x = 120
    for title, body, color in columns:
        draw.rounded_rectangle((x, 220, x + 420, 650), radius=28, fill=(248, 250, 252), outline=pil_rgb(color), width=4)
        draw.rounded_rectangle((x, 220, x + 420, 305), radius=28, fill=pil_rgb(color), outline=pil_rgb(color))
        draw_centered_text(draw, (x + 20, 232, x + 400, 292), title, f(30), "FFFFFF")
        draw_centered_text(draw, (x + 36, 340, x + 384, 610), body, f(25), INK, spacing=12)
        x += 520

    draw_arrow(draw, (540, 435), (635, 435), NAVY, 5)
    draw_arrow(draw, (1060, 435), (1155, 435), NAVY, 5)

    metrics = [
        ("本地测试基线", "1468 passed / 0 failed"),
        ("真实项目 gauntlet", "15 / 15 passed"),
        ("最新状态日期", "2026-05-21"),
    ]
    x = 165
    for label, value in metrics:
        draw.rounded_rectangle((x, 720, x + 360, 795), radius=18, fill=pil_rgb(LIGHT), outline=pil_rgb(BORDER))
        draw_centered_text(draw, (x + 16, 730, x + 344, 758), label, f(20), MUTED)
        draw_centered_text(draw, (x + 16, 758, x + 344, 790), value, f(22), DARK_BLUE)
        x += 455

    img.save(path, quality=95)
    return path


def create_assets() -> dict[str, Path]:
    return {
        "cover": save_cover_image(),
        "positioning": save_positioning_image(),
        "loop": save_loop_image(),
        "architecture": save_architecture_image(),
        "roadmap": save_roadmap_image(),
        "progress": save_progress_image(),
    }


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc: Document, text: str = "", *, style: str | None = None, bold: bool = False, color: str | None = None, size: float | None = None, align=None, after: float | None = None, before: float | None = None, italic: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(run, bold=True)


def add_bullets(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        run = p.add_run(item)
        set_run_font(run)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 10.5, align=None) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def set_table_widths(table, widths_inches: Sequence[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "F4F6F9")
        set_cell_text(cell, header, bold=True, color=DARK_BLUE, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, color=INK, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, title: str, body: str, fill: str = "F4F6F9", accent: str = TEAL) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [0.18, 6.25])
    shade_cell(table.cell(0, 0), accent)
    set_cell_text(table.cell(0, 0), "", size=1)
    shade_cell(table.cell(0, 1), fill)
    cell = table.cell(0, 1)
    cell.text = ""
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=12, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(0)
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption)
    set_run_font(r, size=9.5, color=MUTED, italic=True)


def set_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header_p.add_run("Priority Studio | 投资合作策划书")
    set_run_font(r, size=9, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run("Confidential discussion draft · 2026-05-21")
    set_run_font(r, size=9, color=MUTED)


def add_cover(doc: Document, assets: dict[str, Path]) -> None:
    add_para(doc, "投资合作策划书", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, color=GOLD, bold=True, before=8, after=10)
    add_para(doc, "Priority Studio", align=WD_ALIGN_PARAGRAPH.CENTER, size=30, color=NAVY, bold=True, after=4)
    add_para(doc, "面向中国市场的桌面 AI 项目工作台", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, color=DARK_BLUE, bold=True, after=8)
    add_para(doc, "让普通用户用中文聊天创建、运行、修改和交付小项目", align=WD_ALIGN_PARAGRAPH.CENTER, size=12.5, color=MUTED, after=16)
    add_figure(doc, assets["cover"], "概念图：聊天前台 + 本地工程执行后台 + 国内多模型路由", width=6.4)
    add_callout(
        doc,
        "一句话定位",
        "面向中国普通用户的可安装桌面应用：不展示代码、不要求命令行，让用户通过中文对话获得可运行、可预览、可迭代的小软件、网站和业务工具。",
        fill="F8FAFC",
        accent=GOLD,
    )
    add_para(doc, "版本：合作伙伴 / 投资人沟通稿    日期：2026-05-21    状态：Discussion Draft", align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5, color=MUTED, after=0)
    doc.add_page_break()


def build_doc() -> None:
    assets = create_assets()
    doc = Document()
    set_doc_defaults(doc)
    set_header_footer(doc)
    add_cover(doc, assets)

    add_heading(doc, "执行摘要", 1)
    add_callout(
        doc,
        "核心判断",
        "中国市场需要的不是又一个面向程序员的命令行 Agent，而是一个把 AI 编程能力封装成普通人能使用的桌面项目生产工具。Priority Studio 的机会在于把模型、项目环境、文件系统、验证和交付流程整合成一条用户可理解的闭环。",
        accent=TEAL,
    )
    add_para(
        doc,
        "Priority Studio 计划基于现有 priority-agent Rust runtime，推出 Windows 和 macOS 可安装桌面应用。前台是聊天、项目预览、进度和交付结果；后台负责项目创建、代码修改、依赖安装、终端执行、自动验证、失败修复、快照回滚和多模型路由。"
    )
    add_para(
        doc,
        "本项目的第一阶段目标不是替代专业 IDE，也不是复刻 Claude Code 的 CLI 体验，而是服务“有项目想法但不会编程”的中国普通用户：个体创业者、中小商家、运营人员、学生、设计师和产品经理。今天这些用户要么很难直接使用 Claude Code、Codex 等海外工具，要么无法顺畅获取 GitHub 上的开源项目，要么只能接触国内厂商提供的模型 API、CLI 或 VS Code 插件。Priority Studio 要补上的正是“开箱即用的 AI 编程软件”这一层。"
    )
    add_para(
        doc,
        "成功标准不是模型回答是否漂亮，而是用户是否能得到可运行、可预览、可继续修改的真实项目。"
    )
    add_table(
        doc,
        ["投资/合作方应记住的三点", "说明"],
        [
            ("入口差异", "从 CLI / IDE 转向可安装桌面应用，默认隐藏代码和终端细节。"),
            ("供给变化", "国内模型供给已经成熟，但普通用户仍缺少完整项目交付体验。"),
            ("技术资产", "现有 runtime 已具备 provider、工具执行、权限、记忆、追踪和验证闭环基础。"),
        ],
        [2.0, 4.3],
    )

    add_heading(doc, "1. 投资叙事：为什么是这个项目", 1)
    add_heading(doc, "1.1 问题：强大的 AI 编程工具没有进入普通用户市场", 2)
    add_para(
        doc,
        "海外 AI 编程工具能力强，但对中国普通用户存在多重门槛：账号和网络访问、支付、英文文档、命令行安装、API Key 配置、项目环境配置以及对代码和报错的理解。国内模型能力正在快速进步，但许多产品仍停留在聊天、开发者命令行或 IDE 插件形态，没有把“做项目”的全流程封装为普通用户可以信任的产品。"
    )
    add_heading(doc, "1.2 市场断层：今天的可选路径都不适合普通用户", 2)
    add_para(
        doc,
        "当前中国用户接触 AI 编程能力大致有四条路径，但它们都没有真正面向普通用户完成产品化。这个断层，是 Priority Studio 最应该强调的机会。"
    )
    add_table(
        doc,
        ["用户今天的选择", "现实障碍", "对本项目的启示"],
        [
            ("海外 Claude Code / Codex", "官方形态主要是终端、IDE 或 GitHub 工作流；中国用户还会遇到访问、账号、支付和 API 可用性门槛。", "不能只复制 CLI，需要做国内可直接安装、可直接使用的桌面入口。"),
            ("GitHub 开源项目", "大量 AI 编程项目托管在 GitHub，安装链路依赖 clone、npm、raw.githubusercontent.com、英文 README 和终端命令；这些对普通用户既不稳定也不友好。", "开源项目本身很难普及到大众，需要把能力封装成软件，而不是让用户自己拉仓库。"),
            ("国内大模型厂商", "多数厂商主要提供模型、API、控制台或云平台，模型能力没有自动变成普通用户可完成项目的应用。", "模型是供应，不是产品。产品价值要落在项目闭环、预览、验证、修复和交付。"),
            ("国内编程工具", "即使已有 Qwen Code、Crush 等工具，也多是 CLI、TUI 或 VS Code 形态，默认用户懂命令行、项目目录和 API Key。", "必须隐藏代码和终端细节，把工程能力做成聊天式桌面工作台。"),
        ],
        [1.6, 2.6, 2.1],
    )
    add_callout(
        doc,
        "关键产品判断",
        "普通用户真正需要的不是“更强的模型列表”，而是一个下载安装后就能开始做项目的软件：内置国内模型接入、项目工作区、运行预览、自动验证、失败修复和导出交付。",
        accent=GOLD,
    )

    add_heading(doc, "1.3 方案：把编程 Agent 做成桌面项目工作台", 2)
    add_para(
        doc,
        "Priority Studio 的产品答案是：让用户只面对项目目标、预览结果和关键确认；让 Agent 在本地工作区中完成工程执行。用户说“帮我做一个咖啡店会员积分工具”，应用应能创建项目、安装依赖、启动预览、检查错误、修复失败，并让用户继续用中文改页面、加功能、导出交付。"
    )
    add_heading(doc, "1.4 为什么现在", 2)
    add_bullets(
        doc,
        [
            "国内大模型 API 和平台生态已经形成多供应商选择，适合做 provider-neutral 的多模型路由。",
            "Claude Code、Codex、opencode 等产品已经教育了开发者市场，但其 CLI / IDE / GitHub 形态没有覆盖中国普通用户。",
            "国内模型厂商正在补齐 Coding、Agent、长上下文和多模态能力，但仍缺少一个把这些能力打包成桌面软件的普通用户入口。",
            "中小商家、个体创业者和轻技术岗位有大量小项目需求，传统外包慢、贵且沟通成本高，自己安装开源项目又不现实。",
            "现有 priority-agent 项目已经有 runtime 基础，可以先做桌面产品化，不必从零搭 Agent 内核。",
        ],
    )
    add_figure(doc, assets["positioning"], "图 1：产品定位从开发者工具转向普通用户项目工作台", width=6.35)

    add_heading(doc, "2. 用户与场景", 1)
    add_para(
        doc,
        "第一阶段建议聚焦“非程序员但有明确项目目标”的用户，而不是与专业 IDE 抢高端开发者。这个选择能让产品叙事更清晰：我们卖的不是代码编辑效率，而是项目从想法到可运行结果的完成率。"
    )
    add_table(
        doc,
        ["用户群体", "高频任务", "为什么会付费"],
        [
            ("个体创业者", "落地页、报名系统、订单小工具、演示原型", "比找外包更快，比自己学编程更低门槛"),
            ("中小商家/运营", "会员工具、活动页、商品展示、表格自动化", "直接解决经营和运营效率问题"),
            ("学生/研究人员", "课程项目、数据可视化、实验工具", "需要能跑的项目，而不是只要代码片段"),
            ("产品经理/设计师", "交互原型、内部工具、需求验证 demo", "把原型推进到可试用版本，减少沟通成本"),
        ],
        [1.35, 2.55, 2.4],
    )
    add_heading(doc, "2.1 首批可验证任务", 2)
    add_bullets(
        doc,
        [
            "个人作品集、产品落地页、活动报名页、客户表单。",
            "本地数据看板、Excel/CSV 处理工具、轻量 CRM、订单统计。",
            "书单/笔记/待办类本地 Web App，支持浏览器预览和导出。",
            "面向小商家的会员积分、优惠券、预约登记等原型。",
        ],
    )

    add_heading(doc, "3. 产品形态与核心体验", 1)
    add_callout(
        doc,
        "产品原则",
        "不要把用户带进代码世界，而是把工程能力带到用户的项目世界。默认隐藏代码，保留高级模式和审计能力。",
        accent=BLUE,
    )
    add_para(
        doc,
        "桌面应用主界面建议采用三栏或两栏结构：项目历史、聊天任务流、运行预览。普通用户只看自然语言、状态、预览和交付物；懂技术的用户可以打开高级模式查看文件、日志、命令和 diff。"
    )
    add_figure(doc, assets["loop"], "图 2：从需求到交付的项目闭环", width=6.35)
    add_heading(doc, "3.1 MVP 体验闭环", 2)
    add_numbered(
        doc,
        [
            "用户输入项目目标，并上传必要素材或文件。",
            "Agent 询问少量关键问题，选择技术栈和模板。",
            "后台创建独立工作区，写入文件并安装依赖。",
            "应用启动本地预览，并把结果展示给用户。",
            "验证失败时自动诊断和修复；修复过程可审计但默认折叠。",
            "用户继续用中文修改，应用保存快照并支持回退。",
            "用户导出 zip、发布静态站或保存项目快照。",
        ],
    )

    add_heading(doc, "4. 技术基础与可复用资产", 1)
    add_para(
        doc,
        "现有 priority-agent 并不是一个简单 CLI 壳，而是已经形成了可复用的本地 Agent runtime。桌面版最有价值的路径，是把这些能力从命令行迁移到后台服务，再用桌面前台重新包装。"
    )
    add_table(
        doc,
        ["资产", "当前价值", "桌面版落点"],
        [
            ("Rust runtime", "适合本地执行、文件系统、终端进程和安全边界", "作为 Tauri 后台核心"),
            ("多 Provider", "已有 Kimi、MiniMax、OpenAI-compatible、Custom 等接入结构", "扩展 DeepSeek、GLM、Qwen、百炼渠道"),
            ("工具执行", "文件读写、搜索、命令执行、验证、恢复元数据", "隐藏代码但保留执行闭环"),
            ("权限/风险", "工具审批、风险信号、恢复计划", "桌面权限弹窗、项目沙箱、快照回滚"),
            ("记忆/追踪", "turn trace、tool record、用户偏好和项目上下文", "项目历史、用户习惯、失败经验沉淀"),
            ("验证体系", "真实项目 gauntlet、required validation、closeout evidence", "用项目完成率作为产品指标"),
        ],
        [1.35, 2.65, 2.25],
    )
    add_figure(doc, assets["architecture"], "图 3：建议技术架构", width=6.35)

    add_heading(doc, "4.1 当前开发进展", 2)
    add_callout(
        doc,
        "进展判断",
        "当前项目已经完成的是 Agent runtime 底座，而不是面向普通用户的桌面产品。这个状态对合作方反而是清晰的：底层工程能力已有证据，下一阶段要把它产品化成开箱即用的桌面应用。",
        accent=TEAL,
    )
    add_para(
        doc,
        "截至 2026-05-21 的项目状态文档记录，Priority Agent 已经从早期桌面 Agent 设想演进为可工作的 Rust 编程 Agent runtime：具备交互式 CLI、意图路由、工具执行、权限、记忆、trace、恢复计划、MCP 健康检查、required validation closeout、多 provider 协议处理和真实项目评测基础。"
    )
    add_table(
        doc,
        ["模块/能力", "当前进展", "对桌面版的价值"],
        [
            ("Agent Runtime", "已形成交互式 coding CLI 和状态化 turn runtime，包含 intent routing、session goal、trace、recovery plan。", "可作为桌面应用后台，不必从零写 Agent 执行循环。"),
            ("工具与文件执行", "已支持文件读写、搜索、bash/terminal、验证命令、工具结果记录和失败恢复元数据。", "支撑后台自动创建项目、修改文件、安装依赖、运行预览。"),
            ("权限与安全", "已有权限规则、风险信号、approval response、permission denial recovery 和 trace 证据。", "可迁移为桌面权限弹窗、项目沙箱、信任策略和回滚机制。"),
            ("记忆与上下文", "已有 memory namespace search、conflict hints、retrieval context、记忆行为断言和产品成熟度测试。", "可沉淀用户偏好、项目历史、常用技术栈和失败修复经验。"),
            ("多模型/provider", "已有 Kimi、MiniMax、OpenAI-compatible、Custom 等 provider 架构，并完成第一批 provider-protocol 回归矩阵。", "可扩展 DeepSeek、GLM、Qwen、百炼等国内模型，形成 provider-neutral 供应策略。"),
            ("验证与评测", "状态文档记录最新 deterministic local tests 为 1468 passed / 0 failed；real-project-coding gauntlet 为 15/15 passed。", "可以用项目完成率和验证证据向合作方证明不是 demo-only。"),
        ],
        [1.3, 2.55, 2.5],
    )
    add_figure(doc, assets["progress"], "图 4：当前开发进展与桌面产品化关系", width=6.35)

    add_heading(doc, "4.2 尚未完成的产品化工作", 2)
    add_para(
        doc,
        "现阶段也要诚实说明边界：priority-agent 的强项在后台 runtime，弱项在普通用户产品体验。后续开发重点不是继续堆 CLI 命令，而是把已有 runtime 包装成桌面软件，并完成安装、预览、导出、模型套餐、用户引导和商业化闭环。"
    )
    add_table(
        doc,
        ["待补齐事项", "为什么重要", "建议优先级"],
        [
            ("Tauri 桌面壳", "把 CLI runtime 变成 Windows/macOS 可安装应用，是普通用户可用的前提。", "P0"),
            ("聊天 + 预览界面", "用户不应看到代码和终端，而应看到项目进度、运行预览和交付结果。", "P0"),
            ("项目工作区与沙箱", "每个项目独立运行，限制误删和越权访问，支持快照与回退。", "P0"),
            ("国内模型配置/套餐", "普通用户不能自己处理复杂 API Key 和 provider 配置，需要内置或半托管方案。", "P0"),
            ("非程序员 onboarding", "要用模板、示例项目和少量问题引导用户完成第一个项目。", "P1"),
            ("内测数据闭环", "需要记录完成率、失败原因、平均成本、用户修改轮次和付费意愿。", "P1"),
        ],
        [1.7, 3.35, 1.0],
    )

    add_heading(doc, "5. 国内模型与供应策略", 1)
    add_para(
        doc,
        "模型必须被视为可替换供应商，而不是产品护城河本身。真正可持续的价值在于：本地 runtime、项目上下文、验证证据、失败修复经验、用户工作流记忆和桌面交付体验。"
    )
    add_table(
        doc,
        ["模型/渠道", "建议角色", "产品策略"],
        [
            ("DeepSeek", "通用推理与性价比", "作为默认通用任务和成本敏感用户选项"),
            ("Kimi", "长上下文、Agent、代码和多模态方向", "用于长项目上下文、复杂改造和资料型任务"),
            ("GLM", "Agentic Coding、多模态 Coding、国产生态合作", "用于 Coding 套餐、企业合作和视觉类项目"),
            ("Qwen / Qwen-Coder", "代码模型、百炼渠道、OpenAI 兼容生态", "用于代码生成、前端项目和企业云渠道"),
            ("MiniMax", "备选供应和多模态能力", "用于冗余、成本优化和特定场景补充"),
        ],
        [1.45, 2.25, 2.65],
    )
    add_heading(doc, "5.1 路由原则", 2)
    add_bullets(
        doc,
        [
            "简单澄清和普通聊天使用低成本模型。",
            "大范围代码生成、复杂修复和验证失败后升级到更强模型。",
            "长项目上下文优先使用长上下文模型或上下文缓存。",
            "前端视觉复刻和图片输入优先使用多模态 Coding 能力。",
            "所有模型调用都记录成本、时延、失败率和任务完成贡献。",
        ],
    )

    add_heading(doc, "6. 商业模式与 GTM", 1)
    add_para(
        doc,
        "早期商业模式不宜过早复杂化。建议先用个人订阅和项目包验证付费意愿，再向团队版、企业私有化和模型渠道合作扩展。"
    )
    add_table(
        doc,
        ["模式", "适合对象", "验证重点"],
        [
            ("免费试用", "首次体验用户", "能否在 10 分钟内看到第一个可运行项目"),
            ("个人订阅", "学生、个体创业者、轻技术用户", "月活、复用频率、模型成本覆盖"),
            ("项目包", "低频但明确需求用户", "单项目完成率和客单价"),
            ("团队版", "小团队、工作室、运营部门", "共享项目、权限和企业模型账号"),
            ("私有化/内网", "数据敏感企业", "本地部署、审计、模型私有化对接"),
        ],
        [1.3, 2.35, 2.65],
    )
    add_heading(doc, "6.1 早期获客建议", 2)
    add_bullets(
        doc,
        [
            "先找 10-20 个真实用户访谈，收集他们愿意为哪些小项目付费。",
            "围绕“做一个可运行项目”而非“AI 编程”做 Demo 视频和案例。",
            "从创业社群、电商运营、产品经理、学生项目和设计师原型场景切入。",
            "用项目完成率、平均完成时间、用户修改轮次和付费意愿来判断方向。",
        ],
    )

    add_heading(doc, "7. 6 个月路线图", 1)
    add_figure(doc, assets["roadmap"], "图 4：6 个月产品化路线图", width=6.35)
    add_heading(doc, "7.1 阶段验收指标", 2)
    add_table(
        doc,
        ["阶段", "目标", "硬指标"],
        [
            ("1. 桌面 POC", "Tauri + Rust runtime 打通", "个人作品集网站能创建、运行、预览、修改一次"),
            ("2. 非程序员 MVP", "小型 Web 项目和自动化工具闭环", "20 个典型任务中至少 14 个完成首版可运行结果"),
            ("3. 内测验证", "验证留存、成本和付费场景", "50-100 名内测用户，找到 2-3 个高频付费场景"),
        ],
        [1.25, 2.6, 2.45],
    )
    add_heading(doc, "7.2 后续开发计划：从 Runtime 到桌面软件", 2)
    add_para(
        doc,
        "后续计划建议采用“先产品闭环、再扩大能力”的顺序。不要先做完整 IDE，也不要先做所有语言和所有模型；先证明普通用户能通过桌面应用完成一个小项目，再扩大任务类型和商业化能力。"
    )
    add_table(
        doc,
        ["时间窗口", "开发重点", "可交付结果"],
        [
            ("0-30 天", "桌面 POC：Tauri 项目骨架、聊天前台、Rust runtime 后台调用、项目工作区。", "用户输入一句中文需求后，能创建并打开一个本地 Web 项目。"),
            ("31-60 天", "项目闭环：运行预览、依赖安装、失败诊断、自动修复、快照回退、导出 zip。", "完成“个人作品集/活动页/本地工具”三类模板闭环。"),
            ("61-120 天", "国内模型和体验：接入 2-3 个国内 provider，做模型路由、成本记录、非程序员 onboarding。", "20 个典型任务中至少 14 个完成首版可运行结果。"),
            ("121-180 天", "内测和商业化：50-100 名内测用户、套餐/额度、失败上报、用户访谈和场景复盘。", "找到 2-3 个高频付费场景，形成下一轮融资/合作证据包。"),
        ],
        [1.25, 3.0, 2.05],
    )

    add_heading(doc, "8. 合作诉求", 1)
    add_para(
        doc,
        "当前最需要的不是泛泛的资源，而是能加速桌面产品化和早期场景验证的合作。建议对外明确寻找以下角色与资源。"
    )
    add_table(
        doc,
        ["合作类型", "具体需求", "合作价值"],
        [
            ("核心合伙/早期团队", "桌面应用工程、前端体验设计、Agent runtime 产品化", "补齐从 CLI runtime 到可安装桌面产品的关键能力"),
            ("模型渠道", "DeepSeek、Moonshot、智谱、百炼、MiniMax 等 API/套餐合作", "降低成本、提高稳定性、形成国内模型集成优势"),
            ("种子资金/天使", "支持 6 个月产品化、内测和模型成本", "换取早期切口和技术资产的增长期权"),
            ("场景合作方", "中小商家、教育、设计、运营等真实任务池", "用真实需求验证项目完成率和付费场景"),
        ],
        [1.35, 2.65, 2.35],
    )
    add_callout(
        doc,
        "对外 Ask 建议",
        "我们正在寻找能够共同完成桌面产品化、模型渠道接入和真实用户场景验证的合作伙伴。资金不是唯一诉求，更重要的是产品、工程、模型和场景资源形成闭环。",
        accent=GOLD,
    )

    add_heading(doc, "9. 风险与应对", 1)
    add_table(
        doc,
        ["风险", "表现", "应对"],
        [
            ("模型不稳定", "项目生成失败、修复循环", "多模型路由、验证闭环、任务测试集、失败归因"),
            ("成本不可控", "一次项目消耗大量 token", "上下文缓存、低成本模型分层、预算提示、工具数量控制"),
            ("桌面环境复杂", "依赖安装失败、预览启动失败", "内置运行环境、模板化技术栈、诊断脚本、失败恢复"),
            ("用户需求发散", "什么都想做导致都做不好", "MVP 聚焦 Web 小项目和自动化工具"),
            ("安全和信任", "误删文件、泄露密钥、执行危险命令", "项目沙箱、权限审批、快照回滚、敏感信息扫描"),
            ("竞品跟进", "大厂推出类似入口", "深耕本地工作流、项目完成率、用户记忆和垂直场景"),
        ],
        [1.45, 2.1, 2.75],
    )

    add_heading(doc, "10. 90 天执行计划", 1)
    add_numbered(
        doc,
        [
            "第 1-2 周：确定桌面技术栈、产品信息架构、首批 20 个非程序员任务。",
            "第 3-5 周：搭建 Tauri POC，打通聊天前台与 priority-agent runtime 后台。",
            "第 6-8 周：完成一个 Web 项目模板闭环，包括创建、预览、修改、验证和导出。",
            "第 9-10 周：接入 2-3 个国内模型 provider，加入成本和失败率记录。",
            "第 11-12 周：邀请 10-20 位目标用户测试，形成第一版完成率和付费反馈报告。",
        ],
    )
    add_callout(
        doc,
        "下一次会议建议讨论",
        "1）是否接受“AI 项目工作台”作为外部定位；2）首个 MVP 场景是否聚焦 Web 小项目；3）合伙/合作资源优先级；4）是否启动 Tauri POC。",
        accent=TEAL,
    )

    doc.add_page_break()
    add_heading(doc, "附录 A：写作框架与公开资料依据", 1)
    add_para(
        doc,
        "本策划书按投资人/合作伙伴材料的阅读习惯重写：先给一句话定位和执行摘要，再说明问题、方案、为什么现在、市场切口、产品、商业模式、路线图、团队/合作诉求和风险。Sequoia 的商业计划框架强调 purpose、problem、solution、why now、market、competition、business model、team、financials 和 vision；YC seed deck 模板强调清晰、简洁和叙事；DocSend/Dropbox 的 pitch deck 资料提醒早期材料需要每页有清晰信息点，并突出产品准备度、商业模式和竞争格局。"
    )
    add_table(
        doc,
        ["资料", "用于本文的判断"],
        [
            ("Sequoia - Writing a Business Plan", "采用 purpose/problem/solution/why now/business model/vision 的核心叙事顺序"),
            ("YC Seed Deck Template", "强化清晰、简洁、叙事，而不是堆砌宏观口号"),
            ("DocSend / Dropbox pitch deck research", "投资人阅读时间有限，材料需要图表化、可快速吸收"),
            ("项目 README / PROJECT_STATUS", "提炼当前 runtime、provider、权限、记忆、验证、gauntlet 和产品化进展"),
            ("OpenAI / Anthropic 官方文档", "Codex CLI、Claude Code 仍以终端、IDE、GitHub 等开发者入口为主，且服务地区存在官方限制"),
            ("GitHub / GreatFire 相关资料", "GitHub 和 raw.githubusercontent.com 访问链路在中国存在不稳定和受限问题，难以作为普通用户分发入口"),
            ("DeepSeek / Kimi / GLM / Qwen / MiniMax 官方文档", "证明国内模型供给具备多供应商基础"),
            ("Qwen Code / Crush 文档", "说明国内编程工具仍主要面向终端/开发者，普通用户桌面入口仍有空间"),
        ],
        [2.25, 4.05],
    )
    add_heading(doc, "附录 B：参考链接", 1)
    sources = [
        "Sequoia Capital: https://sequoiacap.com/article/writing-a-business-plan/",
        "Y Combinator seed deck: https://yc.relayto.com/e/yc-seed-qe6drohf",
        "DocSend Startup Index: https://www.docsend.com/pitch-deck-metrics/",
        "Dropbox / DocSend pitch deck research: https://www.dropbox.com/en_GB/resources/docsend-pitch-deck-research",
        "项目当前状态: docs/PROJECT_STATUS.md",
        "项目 README: README.md",
        "真实项目评测计划: docs/REAL_PROJECT_CODING_GAUNTLET_PLAN_2026-05-17.md",
        "产品原则: docs/PERSONAL_AGENT_PRODUCT_PRINCIPLES_2026-05-18.md",
        "OpenAI supported countries / unsupported country help: https://help.openai.com/en/articles/8983035-which-countries-does-openai-currently-support-chatgpt-voice-mode%23.iso",
        "OpenAI Codex CLI getting started: https://help.openai.com/en/articles/11096431-openai-codex-cli-getting-tarted",
        "Claude Code first-day guide: https://support.claude.com/en/articles/14552382-your-first-day-in-claude-code",
        "Anthropic supported regions: https://docs.claude.com/zh-CN/api/supported-regions",
        "GreatFire Analyzer overview: https://en.greatfire.org/analyzer",
        "GreatFire archived raw.githubusercontent.com test: https://archive.is/cF6An",
        "DeepSeek API docs: https://api-docs.deepseek.com/api/list-models",
        "Kimi API docs: https://platform.kimi.ai/docs/models",
        "智谱 GLM 模型总览: https://docs.bigmodel.cn/cn/guide/start/model-overview",
        "阿里云百炼模型列表: https://help.aliyun.com/zh/model-studio/models",
        "Qwen Code: https://help.aliyun.com/zh/model-studio/qwen-code",
        "智谱 Crush: https://docs.bigmodel.cn/cn/guide/develop/crush",
        "Qwen-Coder: https://help.aliyun.com/zh/model-studio/qwen-coder",
    ]
    add_bullets(doc, sources)
    add_para(
        doc,
        "声明：本文为早期合作沟通稿，未包含财务预测、融资条款或正式市场规模测算。下一版如用于正式融资，应补充创始团队履历、产品 Demo 截图、用户访谈证据、模型成本测算和可引用市场规模数据。",
        size=10,
        color=MUTED,
        italic=True,
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
